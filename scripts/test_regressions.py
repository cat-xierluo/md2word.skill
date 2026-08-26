#!/usr/bin/env python3
"""md2word 已知出版逃逸的回归测试。"""

from pathlib import Path
from tempfile import TemporaryDirectory
import sys
import unittest
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image


HERE = Path(__file__).resolve().parent
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))

from formatter import convert_quotes_to_chinese, parse_text_formatting  # noqa: E402
from footnote_handler import (  # noqa: E402
    FootnoteManager,
    _footnote_text_to_runs_xml,
    _inject_footnotes_into_docx,
)
from table_handler import _calc_column_widths, contains_markdown_formatting  # noqa: E402

import md2word  # noqa: E402


class Md2WordRegressionTest(unittest.TestCase):
    def test_plain_identifiers_keep_intraword_underscores_without_emphasis(self):
        config = md2word.get_preset("book-publish")
        md2word.set_config(config)
        paragraph = Document().add_paragraph()
        identifiers = [
            "payment_instance_id",
            "dispute_amount_band",
            "manual_review_required",
            "audit_log_summary",
            "API_SERVER_KEY",
            "main_chart_type",
            "matter_id",
            "foo__bar__baz",
            "rule_2026_version",
            "合规_风险_等级",
        ]
        parse_text_formatting(
            paragraph,
            "、".join(identifiers)
            + "；_正常斜体_、__正常粗体__、___正常粗斜体___。",
        )

        for identifier in identifiers:
            self.assertIn(identifier, paragraph.text)
            identifier_runs = [run for run in paragraph.runs if identifier in run.text]
            self.assertEqual(
                len(identifier_runs),
                1,
                f"技术标识 {identifier} 应作为一个普通文本 run 原样保留",
            )
            self.assertFalse(bool(identifier_runs[0].italic))
            self.assertFalse(bool(identifier_runs[0].bold))
            self.assertIsNone(identifier_runs[0]._r.rPr.find(qn("w:i")))

        formatted_runs = {
            run.text: run
            for run in paragraph.runs
            if run.text in {"正常斜体", "正常粗体", "正常粗斜体"}
        }
        self.assertEqual(set(formatted_runs), {"正常斜体", "正常粗体", "正常粗斜体"})
        self.assertTrue(formatted_runs["正常斜体"].italic)
        self.assertFalse(bool(formatted_runs["正常斜体"].bold))
        self.assertTrue(formatted_runs["正常粗体"].bold)
        self.assertFalse(bool(formatted_runs["正常粗体"].italic))
        self.assertTrue(formatted_runs["正常粗斜体"].bold)
        self.assertTrue(formatted_runs["正常粗斜体"].italic)

    def test_markdown_table_keeps_intraword_underscores_without_emphasis(self):
        identifiers = [
            "payment_instance_id",
            "dispute_amount_band",
            "manual_review_required",
            "audit_log_summary",
            "API_SERVER_KEY",
            "main_chart_type",
            "matter_id",
            "foo__bar__baz",
            "rule_2026_version",
            "合规_风险_等级",
        ]
        for identifier in identifiers:
            self.assertFalse(
                contains_markdown_formatting(identifier),
                f"技术标识 {identifier} 不应触发表格 Markdown 格式分派",
            )
        for formatted in ("_正常斜体_", "__正常粗体__", "___正常粗斜体___"):
            self.assertTrue(contains_markdown_formatting(formatted))

        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "intraword-underscore-table.md"
            output = temp_dir / "intraword-underscore-table.docx"
            markdown.write_text(
                "# 表格下划线回归\n\n"
                "| 字段 | 格式样例 |\n"
                "| --- | --- |\n"
                + "".join(
                    f"| {identifier} | _正常斜体_ / __正常粗体__ / ___正常粗斜体___ |\n"
                    for identifier in identifiers
                ),
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)
            md2word.create_word_document(str(markdown), str(output), config=config)

            table = Document(output).tables[0]
            self.assertEqual(
                [row.cells[0].text for row in table.rows[1:]],
                identifiers,
            )
            for row in table.rows[1:]:
                identifier = row.cells[0].text
                identifier_runs = [
                    run
                    for paragraph in row.cells[0].paragraphs
                    for run in paragraph.runs
                    if identifier in run.text
                ]
                self.assertEqual(len(identifier_runs), 1)
                self.assertFalse(bool(identifier_runs[0].italic))
                self.assertFalse(bool(identifier_runs[0].bold))
                self.assertIsNone(identifier_runs[0]._r.rPr.find(qn("w:i")))

                formatted_runs = {
                    run.text: run
                    for paragraph in row.cells[1].paragraphs
                    for run in paragraph.runs
                    if run.text in {"正常斜体", "正常粗体", "正常粗斜体"}
                }
                self.assertEqual(
                    set(formatted_runs), {"正常斜体", "正常粗体", "正常粗斜体"}
                )
                self.assertTrue(formatted_runs["正常斜体"].italic)
                self.assertTrue(formatted_runs["正常粗体"].bold)
                self.assertTrue(formatted_runs["正常粗斜体"].bold)
                self.assertTrue(formatted_runs["正常粗斜体"].italic)

    def test_code_block_keeps_internal_rhythm_and_adds_configured_outer_spacing(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "code-spacing.md"
            output = temp_dir / "code-spacing.docx"
            markdown.write_text(
                "```text\n"
                "first_line\n"
                "middle_line\n"
                "last_line\n"
                "```\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)
            md2word.create_word_document(str(markdown), str(output), config=config)

            document = Document(output)
            self.assertEqual(len(document.tables), 0)
            paragraphs = {
                paragraph.text: paragraph
                for paragraph in document.paragraphs
                if paragraph.text in {"first_line", "middle_line", "last_line"}
            }
            self.assertEqual(set(paragraphs), {"first_line", "middle_line", "last_line"})
            first, middle, last = (paragraphs[name] for name in ("first_line", "middle_line", "last_line"))
            self.assertEqual(first.paragraph_format.space_before.pt, 6)
            self.assertEqual(first.paragraph_format.space_after.pt, 0)
            self.assertEqual(middle.paragraph_format.space_before.pt, 0)
            self.assertEqual(middle.paragraph_format.space_after.pt, 0)
            self.assertEqual(last.paragraph_format.space_before.pt, 0)
            self.assertEqual(last.paragraph_format.space_after.pt, 6)
            self.assertEqual(first.paragraph_format.line_spacing, 1.2)
            self.assertEqual(first.runs[0].font.size.pt, 9)
            code_fonts = first.runs[0]._element.rPr.rFonts
            self.assertEqual(code_fonts.get(qn("w:ascii")), "Courier New")
            shading = first._p.pPr.find(qn("w:shd"))
            self.assertEqual(shading.get(qn("w:fill")), "F5F5F5")

    def test_six_long_header_columns_fit_usable_page_width(self):
        config = md2word.get_preset("book-publish")
        min_needed_cm = [2.76, 2.76, 2.76, 2.76, 3.40, 3.40]
        cell_lengths = {
            0: [12, 18, 20],
            1: [12, 16, 22],
            2: [12, 20, 24],
            3: [12, 18, 26],
            4: [16, 24, 30],
            5: [16, 22, 28],
        }

        widths = _calc_column_widths(
            cell_lengths,
            6,
            config,
            min_needed_cm=min_needed_cm,
        )
        usable_cm = (
            config.get("page.width")
            - config.get("page.margin_left")
            - config.get("page.margin_right")
        )

        self.assertLessEqual(sum(widths), usable_cm + 1e-9)
        self.assertAlmostEqual(sum(widths), usable_cm, places=6)
        self.assertTrue(all(width >= 1.2 for width in widths))

    def test_markdown_table_ooxml_widths_share_one_fixed_budget(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "wide-table.md"
            output = temp_dir / "wide-table.docx"
            markdown.write_text(
                "# 宽表回归\n\n"
                "| 适用业务类型 | 输入材料要求 | 事实提取方法 | 法律检索策略 | 风险分级与处置 | 输出成果与复核 |\n"
                "| --- | --- | --- | --- | --- | --- |\n"
                "| 合同审查 | 合同及附件 | 逐项抽取事实 | 法规案例并行 | 高中低三级风险 | 修订稿与审查意见 |\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_word_document(str(markdown), str(output), config=config)

            document = Document(output)
            table = document.tables[0]
            tbl_width = int(table._tbl.tblPr.find(qn("w:tblW")).get(qn("w:w")))
            grid_widths = [
                int(grid_col.get(qn("w:w")))
                for grid_col in table._tbl.tblGrid.findall(qn("w:gridCol"))
            ]
            usable_twips = int(
                (
                    config.get("page.width")
                    - config.get("page.margin_left")
                    - config.get("page.margin_right")
                )
                * 1440
                / 2.54
            )

            self.assertEqual(table._tbl.tblPr.find(qn("w:tblLayout")).get(qn("w:type")), "fixed")
            self.assertLessEqual(sum(grid_widths), usable_twips)
            self.assertEqual(tbl_width, sum(grid_widths))
            for row in table.rows:
                cell_widths = [
                    int(cell._tc.tcPr.find(qn("w:tcW")).get(qn("w:w")))
                    for cell in row.cells
                ]
                self.assertEqual(cell_widths, grid_widths)

    def test_plain_markdown_table_caption_is_centered_without_indent(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "table-caption.md"
            output = temp_dir / "table-caption.docx"
            markdown.write_text(
                "# 表题回归\n\n"
                "**表 10-5：鉴定式案例分析步骤**\n\n"
                "| 步骤 | 说明 |\n| --- | --- |\n| 一 | 识别请求权基础 |\n\n"
                '<div align="center">**表10-6：显式居中兼容**</div>\n\n'
                '<div align="center">普通居中文字</div>\n',
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_word_document(str(markdown), str(output), config=config)

            document = Document(output)
            caption = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.startswith("表 10-5：")
            )
            self.assertEqual(caption.alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)
            self.assertEqual(caption.paragraph_format.first_line_indent.pt, 0)
            self.assertEqual(caption.paragraph_format.left_indent.pt, 0)
            self.assertTrue(any(run.bold for run in caption.runs))
            self.assertEqual(caption.runs[0].font.size.pt, 12)

            explicit_caption = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.startswith("表10-6：")
            )
            self.assertEqual(
                explicit_caption.alignment, WD_PARAGRAPH_ALIGNMENT.CENTER
            )
            self.assertEqual(
                explicit_caption.paragraph_format.first_line_indent.pt, 0
            )
            self.assertEqual(explicit_caption.paragraph_format.left_indent.pt, 0)
            self.assertTrue(any(run.bold for run in explicit_caption.runs))
            self.assertEqual(explicit_caption.runs[0].font.size.pt, 12)

            generic_centered = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text == "普通居中文字"
            )
            self.assertEqual(
                generic_centered.alignment, WD_PARAGRAPH_ALIGNMENT.CENTER
            )
            self.assertEqual(
                generic_centered.paragraph_format.first_line_indent.pt, 24
            )

    def test_quote_callouts_share_one_full_width_paragraph_style_without_tables(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = HERE / "fixtures" / "quote-callout.md"
            output = temp_dir / "quote-callout.docx"
            config = md2word.get_preset("book-publish")
            for preset_name in ("book-publish", "legal"):
                preset = md2word.get_preset(preset_name)
                self.assertEqual(
                    preset.get("quote.background_color"),
                    preset.get("code_block.content.background_color"),
                    f"{preset_name} 的引用框与代码框必须使用完全相同的背景色 token",
                )
                self.assertEqual(preset.get("quote.background_color"), "#F5F5F5")
            md2word.set_config(config)

            md2word.create_word_document(str(markdown), str(output), config=config)

            with zipfile.ZipFile(output) as archive:
                document_root = ET.fromstring(archive.read("word/document.xml"))
                document_xml = archive.read("word/document.xml").decode("utf-8")
                footnotes_root = ET.fromstring(archive.read("word/footnotes.xml"))

            self.assertNotIn("[^guide]", document_xml)
            references = list(document_root.iter(qn("w:footnoteReference")))
            self.assertEqual(len(references), 1)
            positive_footnote_texts = [
                "".join(node.text or "" for node in footnote.iter(qn("w:t"))).strip()
                for footnote in footnotes_root.iter(qn("w:footnote"))
                if int(footnote.get(qn("w:id"))) > 0
            ]
            self.assertEqual(positive_footnote_texts, ["导读脚注定义。"])

            document = Document(output)
            self.assertEqual(len(document.tables), 0, "引用块不得再产生 Word 表格或网格线")
            self.assertEqual(len(list(document_root.iter(qn("w:tbl")))), 0)
            exact_six_point_blank_paragraphs = []
            for paragraph in document.paragraphs:
                spacing = paragraph._p.find(f"{qn('w:pPr')}/{qn('w:spacing')}")
                if (
                    not paragraph.text
                    and spacing is not None
                    and spacing.get(qn("w:line")) == "120"
                    and spacing.get(qn("w:lineRule")) == "exact"
                ):
                    exact_six_point_blank_paragraphs.append(paragraph)
            self.assertEqual(
                len(exact_six_point_blank_paragraphs),
                2,
                "案例的两处内部空行各生成一个 exact 6pt 灰底 spacer；连续空行折叠",
            )
            guide = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.startswith("本章导读")
            )
            case_paragraphs = [
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.startswith(("案例：", "第一段说明", "第二段说明"))
            ]
            self.assertEqual(len(case_paragraphs), 3)

            quote_paragraphs = []
            for paragraph in document.paragraphs:
                p_pr = paragraph._p.pPr
                if p_pr is None:
                    continue
                shading = p_pr.find(qn("w:shd"))
                borders = p_pr.find(qn("w:pBdr"))
                if (
                    shading is not None
                    and shading.get(qn("w:fill")) == "F5F5F5"
                    and borders is not None
                    and all(
                        border.get(qn("w:color")) == "F5F5F5"
                        for border in borders
                    )
                ):
                    quote_paragraphs.append(paragraph)
            self.assertEqual(len(quote_paragraphs), 6)
            case_group = quote_paragraphs[1:]
            self.assertEqual(
                [paragraph.text for paragraph in case_group],
                [
                    "案例：统一引用框",
                    "",
                    "第一段说明案例背景，并保留关键判断。",
                    "",
                    "第二段说明处理结果。",
                ],
            )
            for current, following in zip(case_group, case_group[1:]):
                self.assertIs(
                    current._p.getnext(),
                    following._p,
                    "案例从首段到末段必须由连续的 shaded paragraphs 构成",
                )
            self.assertEqual(
                [paragraph._p for paragraph in case_group if not paragraph.text],
                [paragraph._p for paragraph in exact_six_point_blank_paragraphs],
            )

            style_snapshots = []
            quote_groups = [[guide], case_group]
            for group in quote_groups:
                for index, paragraph in enumerate(group):
                    p_pr = paragraph._p.pPr
                    p_pr_tags = [child.tag.rsplit("}", 1)[-1] for child in p_pr]
                    self.assertLess(p_pr_tags.index("pBdr"), p_pr_tags.index("shd"))
                    self.assertLess(p_pr_tags.index("shd"), p_pr_tags.index("spacing"))
                    self.assertLess(p_pr_tags.index("spacing"), p_pr_tags.index("ind"))
                    shading = p_pr.find(qn("w:shd"))
                    self.assertEqual(shading.get(qn("w:fill")), "F5F5F5")
                    borders = p_pr.find(qn("w:pBdr"))
                    self.assertIsNotNone(borders)
                    border_snapshot = {
                        child.tag.rsplit("}", 1)[-1]: (
                            child.get(qn("w:val")),
                            child.get(qn("w:sz")),
                            child.get(qn("w:space")),
                            child.get(qn("w:color")),
                        )
                        for child in borders
                    }
                    expected_edges = {
                        "left": ("single", "2", "6", "F5F5F5"),
                        "right": ("single", "2", "6", "F5F5F5"),
                    }
                    if index == 0:
                        expected_edges["top"] = ("single", "2", "5", "F5F5F5")
                    if index == len(group) - 1:
                        expected_edges["bottom"] = ("single", "2", "5", "F5F5F5")
                    self.assertEqual(
                        border_snapshot,
                        expected_edges,
                        "首段独占 top、末段独占 bottom，中段不得重复累计垂直 padding",
                    )
                    indent = p_pr.find(qn("w:ind"))
                    self.assertIsNone(indent.get(qn("w:left")))
                    self.assertIsNone(indent.get(qn("w:right")))
                    self.assertEqual(paragraph.paragraph_format.first_line_indent.pt, 0)
                    spacing = p_pr.find(qn("w:spacing"))
                    if paragraph.text:
                        self.assertEqual(spacing.get(qn("w:line")), "360")
                        self.assertEqual(spacing.get(qn("w:lineRule")), "auto")
                        self.assertEqual(paragraph.paragraph_format.line_spacing, 1.5)
                        self.assertEqual(paragraph.alignment, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
                    else:
                        self.assertEqual(spacing.get(qn("w:before")), "0")
                        self.assertEqual(spacing.get(qn("w:after")), "0")
                        self.assertEqual(spacing.get(qn("w:line")), "120")
                        self.assertEqual(spacing.get(qn("w:lineRule")), "exact")
                        self.assertNotIn("top", border_snapshot)
                        self.assertNotIn("bottom", border_snapshot)
                    for run in paragraph.runs:
                        if run.text:
                            self.assertEqual(run.font.size.pt, 12)
                    style_snapshots.append(shading.get(qn("w:fill")))

            self.assertEqual(set(style_snapshots), {"F5F5F5"})
            self.assertTrue(any(run.text == "本章导读" and run.bold for run in guide.runs))
            self.assertTrue(any(run.text == "案例：统一引用框" and run.bold for run in case_paragraphs[0].runs))
            self.assertTrue(any(run.text == "关键判断" and run.bold for run in case_paragraphs[1].runs))
            self.assertEqual(len(case_paragraphs), 3)
            self.assertEqual(guide.paragraph_format.space_before.pt, 6)
            self.assertEqual(guide.paragraph_format.space_after.pt, 6)
            self.assertEqual(case_group[0].paragraph_format.space_before.pt, 6)
            self.assertEqual(case_group[-1].paragraph_format.space_after.pt, 6)
            for paragraph in case_group[1:-1]:
                self.assertEqual(paragraph.paragraph_format.space_before.pt, 0)
                self.assertEqual(paragraph.paragraph_format.space_after.pt, 0)
            self.assertEqual(case_group[0].paragraph_format.space_after.pt, 0)
            self.assertEqual(case_group[-1].paragraph_format.space_before.pt, 0)

            ordinary = next(
                paragraph for paragraph in document.paragraphs
                if paragraph.text == "案例后的普通正文。"
            )
            self.assertEqual(ordinary.paragraph_format.first_line_indent.pt, 24)
            self.assertEqual(ordinary.paragraph_format.line_spacing, 1.5)
            self.assertEqual(ordinary.runs[0].font.size.pt, 12)

    def test_quote_list_keeps_marker_and_footnote_inside_callout(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "quote-footnote-list.md"
            output = temp_dir / "quote-footnote-list.docx"
            markdown.write_text(
                "# 引用块脚注与列表回归\n\n"
                "> **本章导读**：三个开源项目。[^chapter-skills]\n"
                "> - 重点条目\n\n"
                "[^chapter-skills]: 项目甲、项目乙与项目丙。\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_word_document(str(markdown), str(output), config=config)

            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                footnotes_root = ET.fromstring(archive.read("word/footnotes.xml"))
            self.assertNotIn("[^chapter-skills]", document_xml)
            positive_footnote_texts = [
                "".join(node.text or "" for node in footnote.iter(qn("w:t"))).strip()
                for footnote in footnotes_root.iter(qn("w:footnote"))
                if int(footnote.get(qn("w:id"))) > 0
            ]
            self.assertEqual(positive_footnote_texts, ["项目甲、项目乙与项目丙。"])

            document = Document(output)
            self.assertEqual(len(document.tables), 0)
            paragraphs = [
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.startswith(("本章导读", "    •"))
            ]
            self.assertEqual(len(paragraphs), 2)
            self.assertTrue(
                any(run.text == "本章导读" and run.bold for run in paragraphs[0].runs)
            )
            self.assertIn("•", paragraphs[1].text)
            self.assertEqual(paragraphs[1].paragraph_format.first_line_indent.pt, 0)

    def test_legacy_quote_cell_margin_maps_to_paragraph_padding(self):
        self.assertEqual(
            md2word._quote_padding_pt(
                {"cell_margin": {"top": 80, "bottom": 120, "left": 140, "right": 160}}
            ),
            {"top": 4.0, "bottom": 6.0, "left": 7.0, "right": 8.0},
        )
        self.assertEqual(
            md2word._quote_padding_pt(
                {
                    "padding": {"top": 3, "bottom": 4, "left": 5, "right": 6},
                    "cell_margin": {"top": 200, "bottom": 200, "left": 200, "right": 200},
                }
            ),
            {"top": 3.0, "bottom": 4.0, "left": 5.0, "right": 6.0},
            "新 padding 必须优先于 v1.3.0 cell_margin 兼容映射",
        )

    def test_data_tables_own_exact_space_after_and_leave_following_content_unchanged(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "table-space-after.md"
            output = temp_dir / "table-space-after.docx"
            image_path = temp_dir / "pixel.png"
            Image.new("RGB", (2, 2), "white").save(image_path)
            markdown.write_text(
                "# 表后留白回归\n\n"
                "| 项目 | 内容 |\n| --- | --- |\n| A | Markdown 表 |\n\n"
                "Markdown 表后的正文。\n\n"
                "<table><tr><th>项目</th><th>内容</th></tr>"
                "<tr><td>B</td><td>HTML 表</td></tr></table>\n\n"
                "HTML 表后的正文。\n\n"
                "![示例图](pixel.png)\n\n"
                "**图 1-1：图注链路保持原样**\n\n"
                "图后的正文。\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)
            md2word.create_word_document(str(markdown), str(output), config=config)

            document = Document(output)
            self.assertEqual(len(document.tables), 2)
            body_children = list(document._element.body)
            table_indexes = [
                index for index, child in enumerate(body_children) if child.tag == qn("w:tbl")
            ]
            self.assertEqual(len(table_indexes), 2)
            for table_index in table_indexes:
                spacer = body_children[table_index + 1]
                self.assertEqual(spacer.tag, qn("w:p"))
                self.assertEqual("".join(node.text or "" for node in spacer.iter(qn("w:t"))), "")
                spacing = spacer.find(f"{qn('w:pPr')}/{qn('w:spacing')}")
                self.assertEqual(spacing.get(qn("w:before")), "0")
                self.assertEqual(spacing.get(qn("w:after")), "0")
                self.assertEqual(spacing.get(qn("w:line")), "120")
                self.assertEqual(spacing.get(qn("w:lineRule")), "exact")

                following = body_children[table_index + 2]
                following_text = "".join(
                    node.text or "" for node in following.iter(qn("w:t"))
                )
                self.assertIn("表后的正文", following_text)
                following_spacing = following.find(f"{qn('w:pPr')}/{qn('w:spacing')}")
                self.assertEqual(following_spacing.get(qn("w:before")), "0")
                self.assertEqual(following_spacing.get(qn("w:after")), "0")
                self.assertEqual(following_spacing.get(qn("w:line")), "360")
                self.assertEqual(following_spacing.get(qn("w:lineRule")), "auto")

            caption = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.startswith("图 1-1：")
            )
            self.assertEqual(caption.paragraph_format.space_before.pt, 3)
            self.assertEqual(caption.paragraph_format.space_after.pt, 8)
            self.assertEqual(caption.paragraph_format.line_spacing, 1.2)
            caption_index = body_children.index(caption._p)
            previous = body_children[caption_index - 1]
            self.assertTrue(
                any(node.tag == qn("w:drawing") for node in previous.iter()),
                "图片与图注之间不得插入数据表专用 exact spacer",
            )

    def test_cjk_ascii_quotes_convert_but_english_apostrophes_survive(self):
        converted = convert_quotes_to_chinese("标注'需律师现场确认'，don't、O'Brien 与 API's 保留。")
        self.assertIn("‘需律师现场确认’", converted)
        self.assertIn("don't", converted)
        self.assertIn("O'Brien", converted)
        self.assertIn("API's", converted)

    def test_footnote_inline_markers_become_word_properties(self):
        xml = _footnote_text_to_runs_xml("*模型概览* 与 **重点**，命令 `book-gate verify`")
        self.assertNotIn("*模型概览*", xml)
        self.assertNotIn("**重点**", xml)
        self.assertNotIn("`book-gate verify`", xml)
        self.assertIn("<w:i/>", xml)
        self.assertIn("<w:b/>", xml)
        self.assertIn('w:ascii="Consolas"', xml)

    def test_injected_footnotes_xml_has_no_literal_markdown(self):
        with TemporaryDirectory() as temp:
            docx_path = Path(temp) / "footnotes.docx"
            Document().save(docx_path)
            _inject_footnotes_into_docx(
                str(docx_path),
                [(1, "*模型概览*"), (2, "**需律师确认**"), (3, "`Skill`")],
            )
            with zipfile.ZipFile(docx_path) as archive:
                xml = archive.read("word/footnotes.xml").decode("utf-8")
            self.assertNotIn("*模型概览*", xml)
            self.assertNotIn("**需律师确认**", xml)
            self.assertNotIn("`Skill`", xml)
            self.assertIn("<w:i/>", xml)
            self.assertIn("<w:b/>", xml)

    def test_endnotes_path_also_removes_markdown_markers(self):
        document = Document()
        manager = FootnoteManager(mode="endnote")
        manager.refs = [(1, "*模型概览* 与 **重点**，命令 `Skill`")]
        manager.append_endnotes_section(document)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("模型概览 与 重点，命令 Skill", text)
        self.assertNotIn("*", text)
        self.assertNotIn("`", text)

    def test_inline_code_markers_are_not_reinterpreted_as_emphasis(self):
        config = md2word.get_preset("book-publish")
        md2word.set_config(config)
        paragraph = Document().add_paragraph()
        parse_text_formatting(
            paragraph,
            "关键词检索（如 `law_keyword`）、语义 / 向量检索（如 `case_vector`）、"
            "精确详情（如 `law_detail`、`case_detail`），运算标识 `op*one` 与 `op*two`，"
            "另有 _下划线斜体_ 和 *星号斜体*。",
        )

        self.assertNotIn("`", paragraph.text)
        self.assertIn("law_keyword", paragraph.text)
        self.assertIn("case_vector", paragraph.text)
        self.assertIn("law_detail", paragraph.text)
        self.assertIn("case_detail", paragraph.text)

        code_names = {
            "law_keyword",
            "case_vector",
            "law_detail",
            "case_detail",
            "op*one",
            "op*two",
        }
        code_runs = [run for run in paragraph.runs if run.text in code_names]
        self.assertEqual({run.text for run in code_runs}, code_names)
        for run in code_runs:
            self.assertFalse(bool(run.italic))
            self.assertEqual(run.font.name, "Courier New")

        italic_runs = {
            run.text: run for run in paragraph.runs if run.text in {"下划线斜体", "星号斜体"}
        }
        self.assertEqual(set(italic_runs), {"下划线斜体", "星号斜体"})
        self.assertTrue(all(run.italic for run in italic_runs.values()))

    def test_repeated_footnote_label_creates_distinct_word_footnotes(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "repeated-footnote.md"
            output = temp_dir / "repeated-footnote.docx"
            markdown.write_text(
                "# 重复脚注\n\n第一次[^same]，第二次[^same]。\n\n[^same]: 相同脚注文本\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_word_document(str(markdown), str(output), config=config)

            with zipfile.ZipFile(output) as archive:
                document_root = ET.fromstring(archive.read("word/document.xml"))
                footnotes_root = ET.fromstring(archive.read("word/footnotes.xml"))
            reference_ids = [
                node.get(qn("w:id"))
                for node in document_root.iter(qn("w:footnoteReference"))
            ]
            footnotes = [
                node
                for node in footnotes_root.iter(qn("w:footnote"))
                if int(node.get(qn("w:id"))) > 0
            ]
            footnote_texts = [
                "".join(node.text or "" for node in footnote.iter(qn("w:t"))).strip()
                for footnote in footnotes
            ]
            self.assertEqual(len(reference_ids), 2)
            self.assertEqual(len(set(reference_ids)), 2, "每次原生脚注引用必须使用独立 w:id")
            self.assertEqual(len(footnotes), 2)
            self.assertEqual(footnote_texts, ["相同脚注文本", "相同脚注文本"])

    def test_adjacent_footnote_references_get_one_superscript_nbsp_separator(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "adjacent-footnotes.md"
            output = temp_dir / "adjacent-footnotes.docx"
            markdown.write_text(
                "# 相邻脚注\n\n"
                "相邻[^a][^b]；空格[^a] [^b]；逗号[^a]，[^b]；顿号[^a]、[^b]。\n\n"
                "[^a]: 脚注甲\n[^b]: 脚注乙\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_word_document(str(markdown), str(output), config=config)

            with zipfile.ZipFile(output) as archive:
                document_root = ET.fromstring(archive.read("word/document.xml"))
            nbsp_runs = []
            reference_paragraph_tokens = None
            for paragraph in document_root.iter(qn("w:p")):
                tokens = []
                reference_count = 0
                for run in paragraph.findall(qn("w:r")):
                    reference = run.find(qn("w:footnoteReference"))
                    if reference is not None:
                        tokens.append(("reference", reference.get(qn("w:id")), run))
                        reference_count += 1
                        continue
                    text = "".join(node.text or "" for node in run.iter(qn("w:t")))
                    if text:
                        tokens.append(("text", text, run))
                    if text == "\u00a0":
                        nbsp_runs.append(run)
                if reference_count == 8:
                    reference_paragraph_tokens = tokens

            self.assertEqual(len(nbsp_runs), 1, "只有源码直接相邻的脚注引用需要 NBSP")
            self.assertIsNotNone(reference_paragraph_tokens)
            nbsp_index = next(
                index
                for index, token in enumerate(reference_paragraph_tokens)
                if token[0] == "text" and token[1] == "\u00a0"
            )
            self.assertEqual(reference_paragraph_tokens[nbsp_index - 1][0], "reference")
            self.assertEqual(reference_paragraph_tokens[nbsp_index + 1][0], "reference")
            run_properties = nbsp_runs[0].find(qn("w:rPr"))
            self.assertEqual(
                run_properties.find(qn("w:vertAlign")).get(qn("w:val")), "superscript"
            )
            self.assertEqual(run_properties.find(qn("w:sz")).get(qn("w:val")), "18")

    def test_positive_footnote_paragraphs_have_compact_explicit_spacing(self):
        with TemporaryDirectory() as temp:
            docx_path = Path(temp) / "footnote-spacing.docx"
            Document().save(docx_path)
            _inject_footnotes_into_docx(
                str(docx_path), [(1, "第一条脚注"), (2, "第二条脚注")]
            )

            with zipfile.ZipFile(docx_path) as archive:
                footnotes_root = ET.fromstring(archive.read("word/footnotes.xml"))
            positive_footnotes = [
                node
                for node in footnotes_root.iter(qn("w:footnote"))
                if int(node.get(qn("w:id"))) > 0
            ]
            self.assertEqual(len(positive_footnotes), 2)
            for footnote in positive_footnotes:
                paragraph = footnote.find(qn("w:p"))
                spacing = paragraph.find(f"{qn('w:pPr')}/{qn('w:spacing')}")
                self.assertIsNotNone(spacing)
                self.assertEqual(spacing.get(qn("w:before")), "0")
                self.assertEqual(spacing.get(qn("w:after")), "0")
                self.assertEqual(spacing.get(qn("w:line")), "240")
                self.assertEqual(spacing.get(qn("w:lineRule")), "auto")

    def test_repeated_endnote_label_reuses_one_number_and_definition(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "repeated-endnote.md"
            output = temp_dir / "repeated-endnote.docx"
            markdown.write_text(
                "# 重复尾注\n\n第一次[^same]，第二次[^same]。\n\n[^same]: 相同尾注文本\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_word_document(
                str(markdown), str(output), config=config, notes_mode="endnote"
            )

            document = Document(output)
            paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
            self.assertIn("第一次1，第二次1。", paragraph_texts)
            self.assertEqual(paragraph_texts.count("[1] 相同尾注文本"), 1)

    def test_external_image_download_function_exists(self):
        # 外链图片下载保持默认启用（原行为），download_external_image 可被直接调用
        self.assertTrue(callable(md2word.download_external_image))
        self.assertFalse(hasattr(md2word, "ALLOW_REMOTE_IMAGES"), "外链图片下载开关已移除，保持默认下载")

    def test_book_mode_keeps_in_chapter_hr_and_only_breaks_between_chapters(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            chapter_one = temp_dir / "ch01.md"
            chapter_two = temp_dir / "ch02.md"
            output = temp_dir / "book.docx"
            chapter_one.write_text(
                "# 第一章\n\n正文一[^note]\n\n---\n\n章内分隔线后的正文。\n\n[^note]: 第一章脚注\n",
                encoding="utf-8",
            )
            chapter_two.write_text(
                "# 第二章\n\n正文二[^note]\n\n[^note]: 第二章脚注\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_book([str(chapter_one), str(chapter_two)], str(output), config)

            document = Document(output)
            horizontal_rule = "─" * config.get("horizontal_rule.repeat_count", 55)
            self.assertEqual(len(document.sections), 2, "两章合并应恰好产生两个 section")
            self.assertEqual(
                sum(paragraph.text == horizontal_rule for paragraph in document.paragraphs),
                1,
                "第一章内的 Markdown 水平线应保留",
            )
            for section in document.sections:
                footnote_properties = section._sectPr.find(qn("w:footnotePr"))
                self.assertIsNotNone(footnote_properties)
                restart = footnote_properties.find(qn("w:numRestart"))
                self.assertEqual(restart.get(qn("w:val")), "eachSec")
            self.assertIn("法律 AI Skill 实战", document.sections[0].header.paragraphs[0].text)
            self.assertTrue(document.sections[1].header.is_linked_to_previous)
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                settings_xml = archive.read("word/settings.xml").decode("utf-8")
            self.assertIn('TOC \\o "1-3" \\h \\z \\u', document_xml)
            self.assertIn("<w:updateFields", settings_xml)

    def test_single_document_first_hr_is_rendered_without_page_break(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "single.md"
            output = temp_dir / "single.docx"
            markdown.write_text(
                "# 单章\n\n首段。\n\n---\n\n第二段。\n\n***\n\n第三段。\n\n___\n\n末段。\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_word_document(str(markdown), str(output), config=config)

            document = Document(output)
            horizontal_rule = "─" * config.get("horizontal_rule.repeat_count", 55)
            self.assertEqual(len(document.sections), 1)
            self.assertEqual(
                sum(paragraph.text == horizontal_rule for paragraph in document.paragraphs),
                3,
                "---、***、___ 都应按 Markdown 语义渲染为水平线",
            )
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertNotIn('<w:br w:type="page"', document_xml)


if __name__ == "__main__":
    unittest.main(verbosity=2)
